from flask import Flask, request, render_template, redirect, url_for, send_file, session, jsonify, flash
import pandas as pd
import google.generativeai as genai
from docx import Document
import os

app = Flask(__name__)
app.secret_key = "your_secret_key"

# Configure the Gemini AI model
GEMINI_API_KEY = "AIzaSyBXD8YgGGsmWcUNQZyzy21DjfaueACQl-g"
genai.configure(api_key=GEMINI_API_KEY)
model = genai.GenerativeModel(model_name="gemini-1.0-pro", generation_config={
    "temperature": 0.7,
    "top_p": 0.95,
    "top_k": 50,
    "max_output_tokens": 1024,
})

# File paths
LOGIN_FILE_PATH = "user_details.xlsx"
FEEDBACK_FILE_PATH = "user_details.xlsx"  # Added feedback file path
MEDICAL_FILE_PATH = r'C:\Users\Jasima Firose\Desktop\my workspace\book 2.xlsx'
HOSPITAL_FILE_PATH = [
        r"C:\Users\Jasima Firose\Desktop\my workspace\Arunachala_Pradesh_hospital_data.csv",
    r"C:\Users\Jasima Firose\Desktop\my workspace\Assam_hospital_data.csv",
    r"C:\Users\Jasima Firose\Desktop\my workspace\Bihar_hospital_data.csv",
    r"C:\Users\Jasima Firose\Desktop\my workspace\Chhattisgarh_data.csv",
    r"C:\Users\Jasima Firose\Desktop\my workspace\Delhi_data.csv",
    r"C:\Users\Jasima Firose\Desktop\my workspace\Goa_hospital_data.csv",
    r"C:\Users\Jasima Firose\Desktop\my workspace\Gujarat_hospital_data.csv",
    r"C:\Users\Jasima Firose\Desktop\my workspace\Haryana_hospital_data.csv",
    r"C:\Users\Jasima Firose\Desktop\my workspace\Himachal_Pradesh_hospital_data.csv",
    r"C:\Users\Jasima Firose\Desktop\my workspace\Jammu_and_Kashmir_data.csv",
    r"C:\Users\Jasima Firose\Desktop\my workspace\Jharkhand_hospital_data.csv",
    r"C:\Users\Jasima Firose\Desktop\my workspace\Karnataka_hospital_data.csv",
    r"C:\Users\Jasima Firose\Desktop\my workspace\Kerala_hospital_data.csv",
    r"C:\Users\Jasima Firose\Desktop\my workspace\Madhya_Pradesh_hospital_data.csv",
    r"C:\Users\Jasima Firose\Desktop\my workspace\Maharashtra_data.csv",
    r"C:\Users\Jasima Firose\Desktop\my workspace\Manipur_hospital_data.csv",
    r"C:\Users\Jasima Firose\Desktop\my workspace\Meghalaya_data.csv",
    r"C:\Users\Jasima Firose\Desktop\my workspace\Mizoram_data.csv",
    r"C:\Users\Jasima Firose\Desktop\my workspace\Nagaland_data.csv",
    r"C:\Users\Jasima Firose\Desktop\my workspace\Odisha_data.csv",
    r"C:\Users\Jasima Firose\Desktop\my workspace\Puducherry_data.csv",
    r"C:\Users\Jasima Firose\Desktop\my workspace\Punjab_data.csv",
    r"C:\Users\Jasima Firose\Desktop\my workspace\Rajasthan_data.csv",
    r"C:\Users\Jasima Firose\Desktop\my workspace\Sikkim_data.csv",
    r"C:\Users\Jasima Firose\Desktop\my workspace\Tamil_Nadu_data.csv",
    r"C:\Users\Jasima Firose\Desktop\my workspace\Telangana_data.csv",
    r"C:\Users\Jasima Firose\Desktop\my workspace\Tripura_data.csv",
    r"C:\Users\Jasima Firose\Desktop\my workspace\Uttar_Pradesh_data.csv",
    r"C:\Users\Jasima Firose\Desktop\my workspace\Uttarakhand_data.csv",
    r"C:\Users\Jasima Firose\Desktop\my workspace\West_Bengal_data.csv"
]

# Load datasets
try:
    medical_data = pd.read_excel(MEDICAL_FILE_PATH)
    hospital_data = pd.concat(
        [pd.read_csv(file, encoding='latin1') for file in HOSPITAL_FILE_PATH],
        ignore_index=True
    )
    hospital_data.fillna('Unknown', inplace=True)
except Exception as e:
    print(f"Error loading files: {e}")
    medical_data, hospital_data = pd.DataFrame(), pd.DataFrame()

@app.route("/get_districts", methods=["POST"])
def get_districts():
    """Return districts based on selected state."""
    state = request.json.get("state")
    if not state:
        return jsonify({"districts": []})

    state = state.strip().lower()
    districts = hospital_data[hospital_data['State'].str.contains(state, na=False, case=False)]
    unique_districts = sorted(districts['District'].dropna().unique().tolist())
    return jsonify({"districts": unique_districts})

@app.route("/", methods=["GET", "POST"])
def login():
    """Login page for collecting user details."""
    if request.method == "POST":
        user_details = {
            "username": request.form.get("username"),
            "age": request.form.get("age"),
            "gender": request.form.get("gender"),
            "email": request.form.get("email"),
            "phone": request.form.get("phone"),
        }
        save_user_details(user_details)
        session.update(user_details)
        return redirect(url_for("index"))
    return render_template("login.html")

@app.route("/app", methods=["GET", "POST"])
def index():
    """Main application logic."""
    chatbot_response = []
    model_prediction = ""
    hospital_results = []
    states = sorted(hospital_data['State'].dropna().unique().tolist())

    # Initialize form values
    user_input = session.get("user_input", "")
    condition = session.get("condition", "")
    state = session.get("state", "")
    district = session.get("district", "")

    if request.method == "POST":
        user_input = request.form.get("user_input", "").strip()
        condition = request.form.get("condition", "").strip()
        state = request.form.get("state", "").strip()
        district = request.form.get("district", "").strip()

        if "predict" in request.form:  # Handle model prediction
            if user_input:
                try:
                    prompt = f"Based on the symptoms described: '{user_input}', predict the most accurate single medical condition like fever, typhoid, etc."
                    chat_session = model.start_chat(history=[{"role": "user", "parts": [prompt]}])
                    model_response = chat_session.send_message(prompt)
                    model_prediction = model_response.text.strip()
                except Exception as e:
                    model_prediction = f"Error processing your input: {e}"

        if "submit" in request.form:  # Handle hospital recommendations
            if user_input:
                chatbot_response = medical_chatbot(user_input)
            if condition:
                chatbot_response.append(f"Condition: {condition}")
                hospital_results = recommend_hospitals(condition, state, district).to_dict(orient='records')

        # Save session data
        session["chatbot_response"] = chatbot_response
        session["hospital_results"] = hospital_results
        session["user_input"] = user_input
        session["condition"] = condition
        session["state"] = state
        session["district"] = district

    return render_template(
        "index.html",
        chatbot_response=chatbot_response,
        model_prediction=model_prediction,
        hospital_results=hospital_results,
        states=states,
        user_input=user_input,
        condition=condition,
        state=state,
        district=district
    )

def save_user_details(details):
    """Save user details to an Excel file."""
    df = pd.DataFrame([details])
    if os.path.exists(LOGIN_FILE_PATH):
        existing_data = pd.read_excel(LOGIN_FILE_PATH)
        updated_data = pd.concat([existing_data, df], ignore_index=True)
        updated_data.to_excel(LOGIN_FILE_PATH, index=False)
    else:
        df.to_excel(LOGIN_FILE_PATH, index=False)

def medical_chatbot(user_input):
    """Chatbot logic to process user input."""
    for _, row in medical_data.iterrows():
        if row['Symptom'].lower() in user_input.lower():
            response = [
                f"Action: {row['Action']}",
                f"Foods to Take: {row['Food_to_taken']}",
                f"Foods to Avoid: {row['Food_to_avoid']}",
                f"Medicine: {row['Medicine']}",
                f"Severity: {row['Severity']}"
            ]
            return response
    return ["I'm sorry, I couldn't understand your problem."]

def recommend_hospitals(medical_condition, state_filter=None, district_filter=None):
    """Recommend hospitals based on the medical condition and filters."""
    state_filter = state_filter.strip().lower() if state_filter else None
    district_filter = district_filter.strip().lower() if district_filter else None
    medical_condition = medical_condition.strip().lower()

    hospital_data['State'] = hospital_data['State'].str.strip().str.lower().fillna('unknown')
    hospital_data['District'] = hospital_data['District'].str.strip().str.lower().fillna('unknown')
    hospital_data['Medical Condition'] = hospital_data['Medical Condition'].str.strip().str.lower()

    filtered_data = hospital_data[hospital_data['Medical Condition'].str.contains(medical_condition, na=False)]
    if state_filter:
        filtered_data = filtered_data[filtered_data['State'] == state_filter]
    if district_filter:
        filtered_data = filtered_data[filtered_data['District'] == district_filter]

    return filtered_data[['Hospital_Name', 'Address', 'Contact_number', 'Rating', 'State', 'District']]

@app.route("/feedback", methods=["POST"])
def feedback():
    """Handle user feedback submissions."""
    feedback_text = request.form.get("feedback", "").strip()
    if feedback_text:
        try:
            # Save feedback to an Excel file
            feedback_data = pd.DataFrame([{"Feedback": feedback_text}])
            if os.path.exists(FEEDBACK_FILE_PATH):
                existing_feedback = pd.read_excel(FEEDBACK_FILE_PATH)
                updated_feedback = pd.concat([existing_feedback, feedback_data], ignore_index=True)
                updated_feedback.to_excel(FEEDBACK_FILE_PATH, index=False)
            else:
                feedback_data.to_excel(FEEDBACK_FILE_PATH, index=False)
            flash("Feedback submitted successfully!", "success")
        except Exception as e:
            flash(f"Error saving feedback: {e}", "danger")
    else:
        flash("Feedback cannot be empty.", "warning")
    return redirect(url_for("index"))

@app.route("/download_receipt", methods=["POST"])
def download_receipt():
    """Generate and download a receipt."""
    user_details = {
        "Username": session.get("username"),
        "Age": session.get("age"),
        "Gender": session.get("gender"),
        "Email": session.get("email"),
        "Phone": session.get("phone"),
    }
    chatbot_response = session.get("chatbot_response", [])
    hospital_results = session.get("hospital_results", [])
    selected_hospitals = request.form.getlist("hospitals")

    # Debugging: Print selected hospitals
    print("Selected Hospitals (From Form):", selected_hospitals)
    print("Hospital Results (From Session):", hospital_results)

    doc = Document()
    doc.add_heading("Medical Chatbot Receipt", level=1)

    # Add user details
    doc.add_heading("User Details", level=2)
    for key, value in user_details.items():
        doc.add_paragraph(f"{key}: {value}")

    # Add chatbot response
    doc.add_heading("Chatbot Response", level=2)
    for line in chatbot_response:
        doc.add_paragraph(f"\u2022 {line}", style="List Bullet")

    # Add selected hospitals
    doc.add_heading("Selected Hospitals", level=2)
    if selected_hospitals:
        for hospital_name in selected_hospitals:
            hospital = next(
                (h for h in hospital_results if h['Hospital_Name'].strip().lower() == hospital_name.strip().lower()),
                None
            )
            if hospital:
                doc.add_paragraph(f"Hospital Name: {hospital['Hospital_Name']}")
                doc.add_paragraph(f"Address: {hospital['Address']}")
                doc.add_paragraph(f"Contact Number: {hospital['Contact_number']}")
                doc.add_paragraph(f"Rating: {hospital['Rating']}")
                doc.add_paragraph(f"State: {hospital['State']}")
                doc.add_paragraph(f"District: {hospital['District']}")
                doc.add_paragraph("")  # Add blank line for spacing
            else:
                doc.add_paragraph(f"Hospital '{hospital_name}' not found in the results.")
    else:
        doc.add_paragraph("No hospitals selected.")

    # Save and send the receipt
    receipt_path = "receipt.docx"
    doc.save(receipt_path)
    return send_file(receipt_path, as_attachment=True)



if __name__ == "__main__":
    app.run(debug=True)
